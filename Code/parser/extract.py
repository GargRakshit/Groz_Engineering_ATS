import io
import os
import re
import shutil
import unicodedata
from pathlib import Path

import pymupdf as fitz
import pytesseract
from docx import Document
from PIL import Image


TESSERACT_CMD = os.getenv("TESSERACT_CMD")
OCR_LANGUAGE = "eng"
OCR_TEXT_THRESHOLD = 100
OCR_RENDER_SCALE = 2


_MOJIBAKE_TRIPLES = {
    "\u2122": "'",     # â€™ → '   (right single quote U+2019; CP1252 0x99 → U+2122)
    "\u02dc": "'",     # â€˜ → '   (left single quote U+2018; CP1252 0x98 → U+02DC)
    "\u0153": '"',     # â€œ → "   (left double quote U+201C; CP1252 0x9C → U+0153)
    "\u009d": '"',     # â€  → "   (right double quote U+201D; CP1252 0x9D unmapped)
    "\u201c": "-",     # â€" → -   (en dash U+2013; CP1252 0x93 → U+201C)
    "\u201d": "-",     # â€" → -   (em dash U+2014; CP1252 0x94 → U+201D)
    "\u00a6": "...",   # â€¦ → ... (ellipsis U+2026; CP1252 0xA6 → U+00A6)
    "\u2039": "",      # â€‹ → ""  (zero-width space U+200B; CP1252 0x8B → U+2039)
    "\u200b": "",      # â€‹ true zero-width variant → drop
}

# Catches any `â€<X>` mojibake triple; resolves common X, drops the rest.
_TRIPLE_RE = re.compile(r"\u00e2\u20ac(.)")
_ZERO_WIDTH = re.compile(r"[\u200b\u200c\u200d\ufeff]")
_NBSP = re.compile(r"[\u00a0\u202f]")
_MULTI_SPACE = re.compile(r"[ \t]{2,}")
# Dingbat and decorative bullet characters that PDFs use as list markers.
# Replaced with "-" so the LLM sees plain ASCII bullets instead of glyphs it
# may echo verbatim into output fields.
_DINGBAT_BULLETS = re.compile(
    r"[\u2022\u2023\u2024\u2027\u2043\u204c\u204d"   # common bullets / asterisms
    r"\u2219\u25aa\u25ab\u25b8\u25cf\u25d8\u25e6"    # geometric bullets
    r"\u2700-\u27bf]"                                 # entire Dingbats block (❑ ✓ ✗ etc.)
)


def _resolve_triple(match):
    return _MOJIBAKE_TRIPLES.get(match.group(1), "")


def clean_extracted_text(text: str) -> str:
    """Remove mojibake / encoding artifacts that PyMuPDF emits from PDFs whose
    text was double-encoded (UTF-8 read as Latin-1). The LLM otherwise copies
    these glyphs verbatim into its JSON output."""
    if not text:
        return text

    text = _TRIPLE_RE.sub(_resolve_triple, text)        # â€<X> mojibake triples
    text = text.replace("\u00c2\u00a0", " ")            # Â + NBSP → space
    text = text.replace("\u00c2", "")                   # lone Â   → drop
    text = text.replace("\u00ef\u00bc", "|")            # ï¼ mojibake (was fullwidth pipe) → |
    text = _ZERO_WIDTH.sub("", text)
    text = _NBSP.sub(" ", text)
    text = _DINGBAT_BULLETS.sub("-", text)              # ❑ ✓ • → plain "-" bullet
    text = unicodedata.normalize("NFKC", text)
    text = _MULTI_SPACE.sub(" ", text)
    return text


def configure_tesseract():
    tesseract_path = Path(TESSERACT_CMD) if TESSERACT_CMD else None

    if tesseract_path and tesseract_path.exists():
        pytesseract.pytesseract.tesseract_cmd = str(tesseract_path)
        return

    if shutil.which("tesseract"):
        return

    raise RuntimeError(
        "Tesseract is not available. Set TESSERACT_CMD to the full path of "
        "tesseract.exe, for example: C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
    )


configure_tesseract()


def normalize_link(url):
    if not url:
        return None

    url = url.strip()

    if url.startswith("mailto:"):
        return url

    if url.startswith("http://") or url.startswith("https://"):
        return url

    if url.startswith("linkedin.com") or url.startswith("github.com") or url.startswith("www."):
        return "https://" + url

    return url


def extract_pdf_text_and_links(path):
    doc = fitz.open(str(path))

    text = "\n".join(page.get_text() for page in doc)

    links = []

    for page_number, page in enumerate(doc, start=1):
        for link in page.get_links():
            raw_url = link.get("uri") or link.get("file")
            url = normalize_link(raw_url)

            if url:
                links.append({
                    "page": page_number,
                    "url": url
                })

    doc.close()

    unique_links = []
    seen = set()

    for link in links:
        url = link["url"]

        if url not in seen:
            unique_links.append(link)
            seen.add(url)

    if len(text.strip()) < OCR_TEXT_THRESHOLD:
        text = extract_pdf_text_with_ocr(path)

    return text, unique_links


def extract_pdf_text_with_ocr(path):
    doc = fitz.open(str(path))
    page_text = []
    matrix = fitz.Matrix(OCR_RENDER_SCALE, OCR_RENDER_SCALE)

    for page_number, page in enumerate(doc, start=1):
        pixmap = page.get_pixmap(matrix=matrix, alpha=False)
        image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        text = pytesseract.image_to_string(image, lang=OCR_LANGUAGE)

        if text.strip():
            page_text.append(f"--- OCR PAGE {page_number} ---\n{text.strip()}")

    doc.close()

    return "\n\n".join(page_text)


def extract_docx_text_and_links(path):
    document = Document(str(path))

    text_parts = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            text_parts.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())

            if row_text:
                text_parts.append(row_text)

    links = []

    for relationship in document.part.rels.values():
        if relationship.reltype.endswith("/hyperlink"):
            url = normalize_link(relationship.target_ref)

            if url:
                links.append({
                    "page": None,
                    "url": url
                })

    unique_links = []
    seen = set()

    for link in links:
        url = link["url"]

        if url not in seen:
            unique_links.append(link)
            seen.add(url)

    return "\n".join(text_parts), unique_links


def _read_txt(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("latin-1", errors="replace")


def extract_document_text_and_links(path):
    extension = path.suffix.lower()

    if extension == ".pdf":
        text, links = extract_pdf_text_and_links(path)
    elif extension == ".docx":
        text, links = extract_docx_text_and_links(path)
    elif extension == ".txt":
        text, links = _read_txt(path), []
    else:
        raise ValueError(f"Unsupported file type: {path.name}")

    return clean_extracted_text(text), links
